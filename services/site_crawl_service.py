import asyncio
import re
import sys
import tempfile
from functools import lru_cache
from typing import Any
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from docx import Document
from pypdf import PdfReader
from rank_bm25 import BM25Okapi

from services.text_chunking_service import chunk_text
from services.token_counter_service import (
    decode_tokens,
    encode_tokens,
    token_count,
)

_DEFAULT_MARKDOWN_GENERATOR_OPTIONS: dict[str, Any] = {
    "ignore_links": True,
    "ignore_images": True,
    "skip_internal_links": True,
    "body_width": 0,
}


@lru_cache(maxsize=1)
def _crawl4ai_stack() -> tuple[Any, Any, Any, Any, Any, Any]:
    """Import crawl4ai only when crawling; avoids heavy DLL init before embedding in MCP."""
    from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig
    from crawl4ai.content_filter_strategy import BM25ContentFilter, PruningContentFilter
    from crawl4ai.markdown_generation_strategy import DefaultMarkdownGenerator

    return (
        AsyncWebCrawler,
        BrowserConfig,
        CrawlerRunConfig,
        BM25ContentFilter,
        PruningContentFilter,
        DefaultMarkdownGenerator,
    )


def _ensure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def _get_markdown_raw(result: Any) -> str:
    md_obj = getattr(result, "markdown", None)

    if md_obj is None:
        return ""

    if isinstance(md_obj, str):
        return md_obj

    return getattr(md_obj, "raw_markdown", "") or ""


def _get_markdown_fit(result: Any) -> str:
    md_obj = getattr(result, "markdown", None)

    if md_obj is None:
        return ""

    if isinstance(md_obj, str):
        return md_obj

    return getattr(md_obj, "fit_markdown", "") or ""


def _get_html(result: Any) -> str:
    return getattr(result, "html", "") or ""


def _truncate_to_max_tokens(
    text: str,
    max_return_tokens: int | None,
    encoding_name: str,
) -> str:
    if max_return_tokens is None:
        return text
    tokens = encode_tokens(text, encoding_name)
    if len(tokens) <= max_return_tokens:
        return text
    return decode_tokens(tokens[:max_return_tokens], encoding_name)


def _pick_markdown_for_chunking(
    markdown_raw: str,
    markdown_fit: str,
    fit_min_chars: int,
) -> tuple[str, str]:
    fit_stripped = markdown_fit.strip()
    if len(fit_stripped) >= fit_min_chars:
        return fit_stripped, "fit"
    return markdown_raw.strip(), "raw"


def _crawler_config_for_fit_markdown(
    *,
    fit_markdown_mode: str,
    user_query: str | None,
    bm25_threshold: float,
    bm25_language: str,
    pruning_threshold: float,
) -> Any:
    _, _, CrawlerRunConfig, BM25ContentFilter, PruningContentFilter, DefaultMarkdownGenerator = (
        _crawl4ai_stack()
    )
    mode = fit_markdown_mode.strip().lower()
    if mode in ("", "off", "none", "raw"):
        return CrawlerRunConfig(verbose=False)
    if mode == "bm25":
        q = (user_query or "").strip()
        if not q:
            return CrawlerRunConfig(verbose=False)
        content_filter = BM25ContentFilter(
            user_query=q,
            bm25_threshold=bm25_threshold,
            language=bm25_language,
        )
    elif mode == "pruning":
        content_filter = PruningContentFilter(threshold=pruning_threshold)
    else:
        raise ValueError(
            "fit_markdown_mode must be 'off', 'bm25', or 'pruning', "
            f"not {fit_markdown_mode!r}"
        )
    return CrawlerRunConfig(
        verbose=False,
        markdown_generator=DefaultMarkdownGenerator(
            content_filter=content_filter,
            options=dict(_DEFAULT_MARKDOWN_GENERATOR_OPTIONS),
        )
    )


def _url_path_suffix(url: str) -> str:
    return urlparse(url).path.lower().rsplit(".", 1)[-1] if "." in urlparse(url).path else ""


def _is_document_url(url: str) -> bool:
    return _url_path_suffix(url) in {"pdf", "docx", "doc"}


def _download_url_bytes(url: str) -> bytes:
    req = Request(
        url,
        headers={
            "User-Agent": "TinySearch/0.1",
            "Accept": "*/*",
        },
    )
    with urlopen(req, timeout=30) as resp:
        return resp.read()


def _extract_pdf_text(data: bytes) -> str:
    with tempfile.SpooledTemporaryFile(max_size=20 * 1024 * 1024) as tmp:
        tmp.write(data)
        tmp.seek(0)
        reader = PdfReader(tmp)
        pages: list[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(f"## Page {idx}\n\n{text}")
        return "\n\n".join(pages).strip()


def _extract_docx_text(data: bytes) -> str:
    with tempfile.SpooledTemporaryFile(max_size=20 * 1024 * 1024) as tmp:
        tmp.write(data)
        tmp.seek(0)
        document = Document(tmp)
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts).strip()


def _extract_document_text(url: str) -> tuple[str, str]:
    suffix = _url_path_suffix(url)
    if suffix == "doc":
        raise ValueError("legacy .doc files are not supported; use PDF or DOCX")

    data = _download_url_bytes(url)
    if suffix == "pdf":
        return _extract_pdf_text(data), "pdf"
    if suffix == "docx":
        return _extract_docx_text(data), "docx"

    raise ValueError(f"unsupported document type: {suffix or 'unknown'}")


def _tokenize_for_bm25(text: str) -> list[str]:
    return re.findall(r"[a-zA-Z0-9_./:#-]+", text.lower())


def rank_chunks_bm25(
    query: str,
    chunks: list[dict],
    top_k: int = 5,
) -> list[dict]:
    if not query or not chunks:
        return []

    corpus = [_tokenize_for_bm25(chunk["text"]) for chunk in chunks]
    bm25 = BM25Okapi(corpus)

    query_tokens = _tokenize_for_bm25(query)
    scores = bm25.get_scores(query_tokens)

    ranked = sorted(
        zip(chunks, scores),
        key=lambda item: item[1],
        reverse=True,
    )

    results = []
    for chunk, score in ranked[:top_k]:
        results.append(
            {
                "chunk_id": chunk["chunk_id"],
                "heading": chunk["heading"],
                "score": float(score),
                "tokens": chunk["tokens"],
                "text": chunk["text"],
            }
        )

    return results


async def crawl(
    url: str,
    max_return_tokens: int = 100000,
    encoding_name: str = "o200k_base",
    *,
    user_query: str | None = None,
    fit_markdown_mode: str = "off",
    fit_min_chars: int = 200,
    bm25_threshold: float = 1.5,
    bm25_language: str = "english",
    pruning_threshold: float = 0.48,
) -> dict:
    """
    Crawl a URL (no internal search).

    Returns url, html, markdown_raw (unfiltered), markdown (chosen for chunking),
    markdown_fit (filtered markdown when a content filter ran, else empty),
    markdown_source ('fit' or 'raw'), tokens_raw.
    """
    _ensure_utf8_stdio()

    AsyncWebCrawler, BrowserConfig, _, _, _, _ = _crawl4ai_stack()

    run_config = _crawler_config_for_fit_markdown(
        fit_markdown_mode=fit_markdown_mode,
        user_query=user_query,
        bm25_threshold=bm25_threshold,
        bm25_language=bm25_language,
        pruning_threshold=pruning_threshold,
    )

    async with AsyncWebCrawler(config=BrowserConfig(verbose=False)) as crawler:
        result = await crawler.arun(url=url, config=run_config)

    html = _get_html(result)
    markdown_raw = _get_markdown_raw(result)
    markdown_fit = (_get_markdown_fit(result) or "") if run_config is not None else ""

    markdown_body, markdown_source = _pick_markdown_for_chunking(
        markdown_raw,
        markdown_fit,
        fit_min_chars,
    )

    markdown_raw = _truncate_to_max_tokens(
        markdown_raw, max_return_tokens, encoding_name
    )
    markdown_body = _truncate_to_max_tokens(
        markdown_body, max_return_tokens, encoding_name
    )
    markdown_fit_out = _truncate_to_max_tokens(
        markdown_fit.strip(), max_return_tokens, encoding_name
    )

    return {
        "url": url,
        "max_return_tokens": max_return_tokens,
        "html": html,
        "markdown_raw": markdown_raw,
        "markdown": markdown_body,
        "markdown_fit": markdown_fit_out if run_config is not None else "",
        "markdown_source": markdown_source,
        "tokens_raw": token_count(markdown_raw, encoding_name),
    }


async def fetch_html_for_query(
    url: str,
    user_query: str,
    *,
    bm25_threshold: float = 1.5,
    bm25_language: str = "english",
) -> dict[str, Any]:
    """Fetch a URL through Crawl4AI applying a BM25 content filter on the body.

    Returns ``final_url``, ``html``, ``markdown_raw``, ``markdown_fit`` and the
    crawler's ``metadata`` dict. The ``final_url`` reflects the URL after any
    redirects Crawl4AI followed.
    """
    _ensure_utf8_stdio()

    AsyncWebCrawler, BrowserConfig, CrawlerRunConfig, BM25ContentFilter, _, DefaultMarkdownGenerator = (
        _crawl4ai_stack()
    )

    bm25_filter = BM25ContentFilter(
        user_query=user_query,
        bm25_threshold=bm25_threshold,
        language=bm25_language,
    )
    config = CrawlerRunConfig(
        verbose=False,
        markdown_generator=DefaultMarkdownGenerator(
            content_filter=bm25_filter,
            options=dict(_DEFAULT_MARKDOWN_GENERATOR_OPTIONS),
        ),
    )

    async with AsyncWebCrawler(config=BrowserConfig(verbose=False)) as crawler:
        result = await crawler.arun(url=url, config=config)

    final_url = (
        getattr(result, "redirected_url", None)
        or getattr(result, "url", None)
        or url
    )
    metadata_obj = getattr(result, "metadata", None)
    metadata = dict(metadata_obj) if isinstance(metadata_obj, dict) else {}

    return {
        "final_url": str(final_url),
        "html": _get_html(result),
        "markdown_raw": _get_markdown_raw(result),
        "markdown_fit": _get_markdown_fit(result) or "",
        "metadata": metadata,
    }


async def crawl_search(
    url: str,
    user_query: str,
    top_k: int = 5,
    max_chunk_tokens: int = 500,
    overlap_tokens: int = 80,
    max_return_tokens: int | None = None,
    encoding_name: str = "o200k_base",
    crawl4ai_bm25_threshold: float = 1.5,
    crawl4ai_language: str = "english",
) -> dict:
    """
    Crawl a URL and run internal "search" over the page.

    - Uses Crawl4AI's BM25ContentFilter to produce `markdown_fit`
    - Chunks markdown (fit-first) and ranks chunks with local BM25

    Returns: url, query, html, markdown_raw, markdown_fit, tokens_*, chunks, ranked_chunks
    """
    _ensure_utf8_stdio()

    if _is_document_url(url):
        markdown_raw, document_type = await asyncio.to_thread(_extract_document_text, url)
        chunks = chunk_text(
            text=markdown_raw,
            max_chunk_tokens=max_chunk_tokens,
            overlap_tokens=overlap_tokens,
            encoding_name=encoding_name,
        )
        ranked_chunks = rank_chunks_bm25(query=user_query, chunks=chunks, top_k=top_k)

        if max_return_tokens is not None:
            for chunk in ranked_chunks:
                tokens = encode_tokens(chunk["text"], encoding_name)
                if len(tokens) > max_return_tokens:
                    chunk["text"] = decode_tokens(tokens[:max_return_tokens], encoding_name)
                    chunk["tokens"] = max_return_tokens

        return {
            "url": url,
            "query": user_query,
            "html": "",
            "markdown_raw": markdown_raw,
            "markdown_fit": markdown_raw,
            "tokens_raw": token_count(markdown_raw, encoding_name),
            "tokens_fit": token_count(markdown_raw, encoding_name),
            "chunks_total": len(chunks),
            "chunks": chunks,
            "ranked_chunks": ranked_chunks,
            "document_type": document_type,
        }

    page = await fetch_html_for_query(
        url=url,
        user_query=user_query,
        bm25_threshold=crawl4ai_bm25_threshold,
        bm25_language=crawl4ai_language,
    )
    html = page["html"]
    markdown_raw = page["markdown_raw"]
    markdown_fit = page["markdown_fit"]
    markdown_for_chunking = markdown_fit or markdown_raw

    chunks = chunk_text(
        text=markdown_for_chunking,
        max_chunk_tokens=max_chunk_tokens,
        overlap_tokens=overlap_tokens,
        encoding_name=encoding_name,
    )

    ranked_chunks = rank_chunks_bm25(query=user_query, chunks=chunks, top_k=top_k)

    if max_return_tokens is not None:
        for chunk in ranked_chunks:
            tokens = encode_tokens(chunk["text"], encoding_name)
            if len(tokens) > max_return_tokens:
                chunk["text"] = decode_tokens(tokens[:max_return_tokens], encoding_name)
                chunk["tokens"] = max_return_tokens

    return {
        "url": url,
        "query": user_query,
        "html": html,
        "markdown_raw": markdown_raw,
        "markdown_fit": markdown_fit,
        "tokens_raw": token_count(markdown_raw, encoding_name),
        "tokens_fit": token_count(markdown_fit, encoding_name) if markdown_fit else 0,
        "chunks_total": len(chunks),
        "chunks": chunks,
        "ranked_chunks": ranked_chunks,
    }


if __name__ == "__main__":
    _ensure_utf8_stdio()

    crawl_result = asyncio.run(
        crawl_search(
            url="https://example.com",
            user_query="What is this page about?",
            top_k=5,
            max_chunk_tokens=500,
            overlap_tokens=80,
        )
    )

    for chunk in crawl_result["ranked_chunks"]:
        print("\n" + "=" * 100)
        print(f"Chunk ID: {chunk['chunk_id']}")
        print(f"Heading: {chunk['heading']}")
        print(f"Score: {chunk['score']}")
        print(f"Tokens: {chunk['tokens']}")
        print("-" * 100)
        print(chunk["text"])
